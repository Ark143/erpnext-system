import os, re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

BRAIN_DIR = r'C:\Users\josem\.gemini\antigravity-ide\brain\ad9d29ca-966c-454f-b5c8-9ae935c95822'
REPO_DOCS = r'c:\Users\josem\erpnext-system\docs'

# Define Screenshot Mapping
SCREENSHOTS = {
    "pos_terminal": os.path.join(BRAIN_DIR, "pos_main_screen_full_1788372031782.png"),
    "pos_catalog_cards": os.path.join(BRAIN_DIR, "pos_catalog_cards_1788373167294.png"),
    "pos_qr_login": os.path.join(BRAIN_DIR, "desk_pos_scan_qr_modal_1788231488356.png"),
    "cashier_badge": os.path.join(BRAIN_DIR, "pos_profile_badge_1788372096745.png"),
    "desk_pos_stock": os.path.join(BRAIN_DIR, "desk_vehicle_pos_in_stock_on_1788335919298.png"),
    "pos_payment_methods": os.path.join(BRAIN_DIR, "pos_payment_methods_notes_1788372062912.png"),
    "pos_history": os.path.join(BRAIN_DIR, "pos_history_screen_1788373176463.png"),
    "workspace_top": os.path.join(BRAIN_DIR, "top_view_charts_kpi_1788339306097.png"),
    "workspace_cards": os.path.join(BRAIN_DIR, "shortcuts_module_cards_1788339339522.png"),
    "vehicle_analytics": os.path.join(BRAIN_DIR, "vehicle_analytics_page_1788339634886.png"),
    "customer_vehicle": os.path.join(BRAIN_DIR, "customer_vehicle_form_1788368486364.png"),
    "vehicle_inspection": os.path.join(BRAIN_DIR, "vehicle_inspection_form_1788368513066.png"),
    "vehicle_job_orders": os.path.join(BRAIN_DIR, "vjo_ultra_mrf_desk_1788337950096.png"),
    "stock_entry_safety": os.path.join(BRAIN_DIR, "stock_entry_receiver_verification_1788368390547.png"),
    "employee_badge": os.path.join(BRAIN_DIR, "employee_qr_badge_1788367919149.png"),
    "pos_invoice_submitted": os.path.join(BRAIN_DIR, "pos_invoice_submitted_1788368405271.png"),
    "purchase_order": os.path.join(BRAIN_DIR, "pur_ord_2026_00010_form_1788337528286.png"),
    "approvals_tab": os.path.join(BRAIN_DIR, "approvals_tab_1788337848963.png"),
    "asset_form": os.path.join(BRAIN_DIR, "asset_form_top_1788312156436.png"),
    "asset_depreciation": os.path.join(BRAIN_DIR, "asset_depreciation_schedule_scrolled_1788312168154.png"),
}

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def add_heading_1(doc, text):
    h = doc.add_heading(level=1)
    h.paragraph_format.space_before = Pt(18)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138) # Navy #1e3a8a
    return h

def add_heading_2(doc, text):
    h = doc.add_heading(level=2)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    run = h.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(37, 99, 235) # Blue #2563eb
    return h

def add_heading_3(doc, text):
    h = doc.add_heading(level=3)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)
    run = h.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)
    return h

def add_p(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.bold = bold
    run.font.italic = italic
    return p

def add_image_box(doc, img_key, caption=""):
    img_path = SCREENSHOTS.get(img_key)
    if not img_path or not os.path.exists(img_path):
        p = doc.add_paragraph(f"[Screenshot Placeholder: {img_key}]")
        p.paragraph_format.space_after = Pt(6)
        return

    # Center-aligned paragraph for the picture
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(img_path, width=Inches(6.0))

    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after = Pt(8)
        crun = cp.add_run(f"Figure: {caption}")
        crun.font.name = 'Calibri'
        crun.font.size = Pt(9.0)
        crun.font.italic = True
        crun.font.color.rgb = RGBColor(100, 116, 139) # Slate 500

def add_callout(doc, text, title="NOTE / TIP"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "EFF6FF") # Light Blue
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"ℹ️ {title}: ")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(30, 58, 138)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def build_visual_guide():
    doc = docx.Document()

    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Document Header
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run("ULTRA MRF — ERPNext & VMS End-to-End Visual User Guide")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(30, 58, 138)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(12)
    r_sub = p_sub.add_run("Step-by-Step Operational Manual with Live Screen Walkthroughs for All Modules")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(13)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(100, 116, 139)

    add_callout(doc, "This manual serves as the standard operational training and user guide for ULTRA MRF branch personnel, cashiers, technicians, warehouse custodians, and accountants. Every process flow is accompanied by actual production screenshots from ERPNext v16 and the Vehicle Management System.", "OFFICIAL OPERATING MANUAL")

    # ==========================================
    # MODULE 1: VEHICLE MANAGEMENT WORKSPACE
    # ==========================================
    add_heading_1(doc, "Module 1: Vehicle Management Workspace & Live Analytics")
    add_p(doc, "The Vehicle Management Workspace is the primary command center in ERPNext Desk. When users navigate to /desk/vehicle-management, the system presents the live Vehicle Analytics & Performance Dashboard as the first view.")
    
    add_heading_2(doc, "1.1 Workspace Top View: KPI Cards & Live Performance Charts")
    add_p(doc, "The top section displays real-time KPI counters (Registered Vehicles, Total Job Orders, POS Invoices, Lifetime Revenue) alongside 4 interactive charts (Job Orders by Company, POS Sales by Company, Customer Vehicles by Make, and Job Orders by Status).")
    add_image_box(doc, "workspace_top", "Vehicle Management Workspace — 1st View Dashboard, KPI Cards & Live Charts")

    add_heading_2(doc, "1.2 Workspace Quick Shortcuts & Module Cards")
    add_p(doc, "Scrolling down reveals the Quick Shortcuts rail with live record count badges (e.g. 38,654 Vehicles, 481 Job Orders, 33 POS Invoices, 570 Bins) and 4 categorized module cards: Operations & Front Desk, Vehicle Masters & Registry, Reports & Analytics, and Inventory & Parts.")
    add_image_box(doc, "workspace_cards", "Vehicle Management Workspace — Quick Shortcuts & Organized Module Cards")

    add_heading_2(doc, "1.3 Interactive Multi-Filter Vehicle Analytics Page")
    add_p(doc, "Clicking the 'Vehicle Analytics Dashboard' shortcut opens the dedicated multi-filter analytics engine at /desk/vehicle_analytics. Users can filter by branch and timespan to inspect labor vs. parts sales splits, top-selling tire models, and popular mechanic services.")
    add_image_box(doc, "vehicle_analytics", "Dedicated Vehicle Analytics Page with Company Filter & Service Breakdown")

    # ==========================================
    # MODULE 2: POINT OF SALE (POS) COUNTER
    # ==========================================
    add_heading_1(doc, "Module 2: Touchscreen Vehicle POS Counter Operations")
    add_p(doc, "The Vehicle POS Terminal provides an ultra-fast checkout experience for retail tire sales, parts, and express bay services. It supports camera QR login, plate number lookups, auto-locked cashier branch assignment, multi-payment options, and dual-ledger ERPNext posting.")

    add_heading_2(doc, "2.1 Cashier Sign-In via QR Badge or Credentials")
    add_p(doc, "Open http://38.247.138.224:10017/pos-terminal. Cashiers can log in by entering their username/password or by scanning their printed employee QR badge using the built-in camera scanner modal.")
    add_image_box(doc, "pos_qr_login", "Web POS Terminal — Camera QR Badge Scanner Modal")

    add_heading_2(doc, "2.2 Unified POS Counter, Product Cards & Auto-Locked Branch")
    add_p(doc, "The POS Counter integrates product browsing (left) and the active Cart/Ticket (right) into one seamless view. On desktop/tablet terminals, having separate tabs for 'Catalog' and 'Ticket' is redundant, so they operate concurrently.")
    add_p(doc, "Product Cards: Each item displays its actual uploaded Item Image (with clean category graphics as fallback), item code, stock level, standardized unit price formatted to support large numbers (up to ₱9,999,999.00), and a standardized +ADD pill button.")
    add_p(doc, "Branch Security: The operating branch is automatically detected from the cashier's Employee record (or Cashier Profile) and displayed as a locked badge (e.g. 'BRANCH: ULTRA MRF'). Cashiers cannot alter the branch via a dropdown, preventing billing cross-contamination.")
    add_image_box(doc, "pos_terminal", "Web POS Terminal — Active Counter Screen with Auto-Locked Branch Badge")
    add_image_box(doc, "pos_catalog_cards", "Product Cards — High-Resolution Item Images, Standardized Prices and +ADD Buttons")
    add_image_box(doc, "desk_pos_stock", "Desk Vehicle POS — Catalog with In-Stock: ON Enabled")

    add_heading_2(doc, "2.3 Payment Method Selection Buttons & Order Notes/Remarks")
    add_p(doc, "Cashiers can instantly select payment tenders with dedicated touch buttons: Cash, Card, GCash, Maya, BDO, and Bank Transfer. Selecting a digital or card payment automatically populates the paid amount to match the bill total.")
    add_p(doc, "A dedicated Notes / Remarks field captures customer instructions, payment reference numbers, check details, and bay notes, saving directly to both Vehicle POS Invoice and ERPNext POS Invoice.")
    add_image_box(doc, "pos_payment_methods", "POS Ticket Panel — Payment Method Quick Buttons & Order Remarks Input")

    add_heading_2(doc, "2.4 Real-Time Transaction History with Live Filters & Sync")
    add_p(doc, "Clicking the Transaction History icon (chart icon) smoothly switches to a full-width view of all recent invoices. Data is fetched in real time from the server upon opening the tab.")
    add_p(doc, "Toolbar Filters: Cashiers and managers can filter invoices by 'All', 'Today', or 'This Month', or select a custom date range (From Date to To Date) with the Apply button. A dedicated '🔄 Refresh' button allows manual on-demand syncing.")
    add_p(doc, "Each invoice card displays the document name, date/time, customer, vehicle plate, company branch, total amount, payment method, payment status (Paid/Draft), and a direct clickable deep link to open the invoice in ERPNext Desk.")
    add_image_box(doc, "pos_history", "POS Terminal — Real-Time Transaction History with Date Filters and Refresh Button")

    add_heading_2(doc, "2.5 Official Cashier ID Profile Card")
    add_p(doc, "Clicking the Cashier Profile icon (user icon) displays a single, official Cashier ID Profile card containing employee name, employee number, designation, assigned company and branch, department, reports-to manager, user email, and scannable vector QR code.")
    add_image_box(doc, "cashier_badge", "POS Terminal — Official Cashier ID Profile Card & Actions")

    add_heading_2(doc, "2.6 Dual-Ledger Financial Posting (POS Invoice)")
    add_p(doc, "When a sale is completed, the system concurrently creates a Vehicle POS Invoice and auto-submits a standard ERPNext POS Invoice (ACC-PSINV-...). This immediately writes General Ledger entries and decrements warehouse stock.")
    add_image_box(doc, "pos_invoice_submitted", "ERPNext Desk — Submitted POS Invoice (ACC-PSINV-2026-00032) with GL Impact")

    # ==========================================
    # MODULE 3: WORKSHOP & VEHICLE OPERATIONS
    # ==========================================
    add_heading_1(doc, "Module 3: Workshop Bay Operations & Job Order Lifecycle")
    add_p(doc, "Automotive workshop operations link vehicle intake, digital inspections, estimating, technician bay dispatch, and parts consumption.")

    add_heading_2(doc, "3.1 Customer Vehicle Registry (38,654 Vehicles)")
    add_p(doc, "Navigate to /desk/customer-vehicle. Search or create customer vehicles. The system indexes plate numbers, makes, models, chassis/VIN, engine numbers, and customer linkages with automatic space normalization.")
    add_image_box(doc, "customer_vehicle", "Customer Vehicle Form — NDB-3344 Master Record and Owner Details")

    add_heading_2(doc, "3.2 Multi-Point Electronic Vehicle Inspection")
    add_p(doc, "Navigate to /desk/vehicle-inspection. Service advisors evaluate vehicle condition upon intake using customizable checklist templates (Tire Tread, Brake Pads, Battery Voltage, Fluid Levels) and attach intake photos.")
    add_image_box(doc, "vehicle_inspection", "Vehicle Inspection Form — Electronic Multi-Point Inspection Checklist")

    add_heading_2(doc, "3.3 Vehicle Job Order Bay Dispatch & Technician Allocation")
    add_p(doc, "Navigate to /desk/vehicle-job-order. Service advisors create job orders tracking labor lines (PMS, Alignment, Balancing) and parts lines (Tires, Oil, Filters). Track status progression: Draft -> In Progress -> Completed -> Released.")
    add_image_box(doc, "vehicle_job_orders", "Vehicle Job Orders List in Desk — Bay Status & Operating Company Breakdown")

    # ==========================================
    # MODULE 4: INVENTORY & MATERIAL HANDOVER
    # ==========================================
    add_heading_1(doc, "Module 4: Inventory Management & Receiver Safety Handover")
    add_p(doc, "Warehouse controls manage 570 bin locations, inter-branch transfers, and mandatory receiver verification to eliminate unauthorized withdrawals.")

    add_heading_2(doc, "4.1 Stock Entry (Material Issue) Safety Verification Check")
    add_p(doc, "When creating a Material Issue (/desk/stock-entry), the 'Receiver Verification & Handover Safety Check' box intercepts submission. Technicians must scan their employee QR badge, and a turnover photo proof must be captured.")
    add_image_box(doc, "stock_entry_safety", "Stock Entry Material Issue — Receiver Verification & Handover Safety Check Box")

    add_heading_2(doc, "4.2 Employee ID Badges with QR Code")
    add_p(doc, "Every employee has an official QR badge printable directly from their Employee record (/desk/employee) or via the POS Terminal. Scanning the QR code instantly verifies the employee's name, ID, and company.")
    add_image_box(doc, "employee_badge", "Official Employee ID Badge with QR Code — Ready to Print for Technicians")

    # ==========================================
    # MODULE 5: PROCUREMENT & APPROVALS
    # ==========================================
    add_heading_1(doc, "Module 5: Buying, Procurement & Executive Approvals")
    add_p(doc, "Governs tire and parts procurement contracts, vendor receipts, and multi-branch management approvals.")

    add_heading_2(doc, "5.1 Purchase Orders & 3-Way Matching")
    add_p(doc, "Navigate to /desk/purchase-order. Issue commercial purchase orders to tire manufacturers and parts distributors. Compares quantities against Goods Receipts and Supplier Invoices.")
    add_image_box(doc, "purchase_order", "Purchase Order Form (PUR-ORD-2026-00010) — Supplier Procurement Details")

    add_heading_2(doc, "5.2 Executive Dashboard Approvals Tab with Direct Desk Routing")
    add_p(doc, "Open /executive-dashboard and click the Approvals tab. Executives view live pending draft cards across 9 DocTypes. Clicking 'View all [DocType]s in Desk ->' opens the canonical desk list view pre-filtered to the operating company.")
    add_image_box(doc, "approvals_tab", "Executive Dashboard — Approvals Tab with Canonical Desk Routing")

    # ==========================================
    # MODULE 6: FIXED ASSETS & DEPRECIATION
    # ==========================================
    add_heading_1(doc, "Module 6: Fixed Assets & Automated Depreciation Accounting")
    add_p(doc, "Capitalizes heavy automotive workshop machinery across all service centers with automated straight-line depreciation accounting under IAS 16.")

    add_heading_2(doc, "6.1 Workshop Machinery Asset Master (10 Capitalized Machines)")
    add_p(doc, "Navigate to /desk/asset. Capitalized machinery includes 3D Wheel Aligners, Automatic Tire Changers, Dynamic Balancers, 2-Post Hydraulic Lifts, and AC Recovery Stations.")
    add_image_box(doc, "asset_form", "Fixed Asset Master Form — 3D HD Wheel Alignment System")

    add_heading_2(doc, "6.2 Automated Monthly Straight-Line Depreciation Schedule")
    add_p(doc, "Under the Depreciation Schedule tab, ERPNext calculates exact monthly depreciation amounts and automatically generates monthly General Ledger journal entries crediting Accumulated Depreciation.")
    add_image_box(doc, "asset_depreciation", "Fixed Asset Form — Straight-Line Monthly Depreciation Schedule Tab")

    # Save to Word files
    out_repo = os.path.join(REPO_DOCS, "ULTRA_MRF_VMS_STEP_BY_STEP_VISUAL_USER_GUIDE.docx")
    out_repo_latest = os.path.join(REPO_DOCS, "ULTRA_MRF_VMS_STEP_BY_STEP_VISUAL_USER_GUIDE_Latest.docx")
    out_brain = os.path.join(BRAIN_DIR, "ULTRA_MRF_VMS_STEP_BY_STEP_VISUAL_USER_GUIDE.docx")

    try:
        doc.save(out_repo)
        print("Saved to:", out_repo)
    except PermissionError:
        doc.save(out_repo_latest)
        print("Note: Primary file open in Word. Saved latest version to:", out_repo_latest)

    doc.save(out_brain)
    print("Saved to artifact directory:", out_brain)

if __name__ == '__main__':
    build_visual_guide()
