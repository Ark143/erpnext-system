import re, os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(30, 41, 59)

    src_path = r'c:\Users\josem\erpnext-system\docs\ERPNEXT_DEPLOYMENT_IMPLEMENTATION_PLAN_ULTRA_MRF.md'
    with open(src_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_buffer = []

    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith('```'):
            if in_code_block:
                in_code_block = False
                code_str = '\n'.join(code_buffer)
                
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = tbl.cell(0, 0)
                set_cell_background(cell, "F1F5F9")
                set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.05
                run = p.add_run(code_str)
                run.font.name = 'Consolas'
                run.font.size = Pt(9.0)
                run.font.color.rgb = RGBColor(15, 23, 42)
                
                sp = doc.add_paragraph()
                sp.paragraph_format.space_after = Pt(4)
                code_buffer = []
            else:
                in_code_block = True
                code_buffer = []
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        if line.strip() in ['---', '***', '___']:
            i += 1
            continue

        if '|' in line and i + 1 < len(lines) and re.match(r'^[\|\s\-\:]+$', lines[i+1].strip()):
            raw_headers = [c.strip() for c in line.split('|')[1:-1]]
            i += 2
            table_rows = []
            while i < len(lines) and '|' in lines[i]:
                row_cells = [c.strip() for c in lines[i].split('|')[1:-1]]
                table_rows.append(row_cells)
                i += 1
            
            col_count = len(raw_headers)
            tbl = doc.add_table(rows=1 + len(table_rows), cols=col_count)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            for c_idx, h_text in enumerate(raw_headers):
                cell = tbl.cell(0, c_idx)
                set_cell_background(cell, "1E3A8A")
                set_cell_margins(cell, top=140, bottom=140, left=160, right=160)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(clean_md_formatting(h_text))
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
            
            for r_idx, r_data in enumerate(table_rows):
                bg_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
                for c_idx in range(col_count):
                    val = r_data[c_idx] if c_idx < len(r_data) else ""
                    cell = tbl.cell(1 + r_idx, c_idx)
                    set_cell_background(cell, bg_color)
                    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)
                    add_formatted_text(p, val)
            
            doc.add_paragraph().paragraph_format.space_after = Pt(6)
            continue

        if line.startswith('# '):
            text = line[2:].strip()
            h = doc.add_heading(level=0)
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(8)
            run = h.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = RGBColor(30, 58, 138)
            i += 1
            continue

        if line.startswith('## '):
            text = line[3:].strip()
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(37, 99, 235)
            i += 1
            continue

        if line.startswith('### '):
            text = line[4:].strip()
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(15, 23, 42)
            i += 1
            continue

        if line.startswith('#### '):
            text = line[5:].strip()
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(2)
            run = h.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(51, 65, 85)
            i += 1
            continue

        if line.strip().startswith('* ') or line.strip().startswith('- '):
            indent_level = (len(line) - len(line.lstrip())) // 2
            bullet_text = re.sub(r'^\s*[\*\-]\s+', '', line)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, bullet_text)
            i += 1
            continue

        m_num = re.match(r'^\s*(\d+)\.\s+(.*)', line)
        if m_num:
            indent_level = (len(line) - len(line.lstrip())) // 2
            num_text = m_num.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(0.25 * (indent_level + 1))
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, num_text)
            i += 1
            continue

        if line.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, line.strip())

        i += 1

    targets = [
        r'c:\Users\josem\erpnext-system\docs\ERPNEXT_DEPLOYMENT_IMPLEMENTATION_PLAN_ULTRA_MRF.docx',
        r'C:\Users\josem\.gemini\antigravity-ide\brain\ad9d29ca-966c-454f-b5c8-9ae935c95822\ERPNEXT_DEPLOYMENT_IMPLEMENTATION_PLAN_ULTRA_MRF.docx'
    ]
    
    saved_paths = []
    for path in targets:
        try:
            doc.save(path)
            saved_paths.append(path)
        except Exception as e:
            print(f"Error saving to {path}: {e}")

    print("Successfully generated Word Document at:")
    for sp in saved_paths:
        print(f" - {sp}")

def clean_md_formatting(s):
    s = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', s)
    s = re.sub(r'\*\*([^*]+)\*\*', r'\1', s)
    s = re.sub(r'\*([^*]+)\*', r'\1', s)
    s = re.sub(r'`([^`]+)`', r'\1', s)
    return s.strip()

def add_formatted_text(paragraph, text):
    tokens = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))', text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**') and len(tok) >= 4:
            run = paragraph.add_run(tok[2:-2])
            run.font.bold = True
        elif tok.startswith('`') and tok.endswith('`') and len(tok) >= 2:
            run = paragraph.add_run(tok[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(194, 65, 12)
        elif tok.startswith('[') and '](' in tok and tok.endswith(')'):
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', tok)
            if m:
                label, url = m.group(1), m.group(2)
                run = paragraph.add_run(label)
                run.font.color.rgb = RGBColor(2, 132, 199)
                run.font.underline = True
            else:
                paragraph.add_run(tok)
        else:
            paragraph.add_run(tok)

if __name__ == '__main__':
    create_document()
